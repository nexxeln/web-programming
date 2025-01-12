from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def read_file_content(file_path):
    """Read and return the content of a file."""
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return None


def add_code_block(doc, code, language):
    """Add a code block with syntax highlighting simulation."""
    # Add the language header
    lang_para = doc.add_paragraph()
    lang_run = lang_para.add_run(f"{language}")
    lang_run.font.size = Pt(10)
    lang_run.font.name = "Consolas"
    lang_run.font.color.rgb = RGBColor(100, 100, 100)

    # Add the code
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.size = Pt(10)
    code_run.font.name = "Consolas"

    # Add spacing after code block
    doc.add_paragraph()


def create_exercise_doc(exercise_num):
    doc = Document()

    # Set up the document title
    title = doc.add_heading(f"Web Programming Lab - Exercise {exercise_num}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add student info
    doc.add_paragraph("Name: Shoubhit Dash")
    doc.add_paragraph("Registration Number: 23BRS1224")
    doc.add_paragraph("Course Code: BCSE203E")
    doc.add_paragraph(f"Live URL: https://webprog.nexxel.dev/{exercise_num}.html")
    doc.add_paragraph()

    # First add the main exercise HTML file
    main_html_path = f"public/{exercise_num}.html"
    if os.path.exists(main_html_path):
        doc.add_heading("Main Exercise Page", level=1)
        doc.add_heading("HTML", level=2)
        content = read_file_content(main_html_path)
        if content:
            add_code_block(doc, content, "HTML")

    # Add the main CSS file
    main_css_path = "src/style.css"
    if os.path.exists(main_css_path):
        doc.add_heading("CSS", level=2)
        content = read_file_content(main_css_path)
        if content:
            add_code_block(doc, content, "CSS")

    # Base directory for the exercise subpages
    base_dir = f"public/exercises/{exercise_num}"

    if not os.path.exists(base_dir):
        print(f"Error: Exercise {exercise_num} directory not found")
        return

    # Get all HTML files in the exercise directory recursively
    html_files = []
    for root, _, files in os.walk(base_dir):
        for file in files:
            if file.endswith(".html"):
                rel_path = os.path.relpath(os.path.join(root, file), base_dir)
                html_files.append(rel_path)

    # Sort files to ensure consistent ordering
    html_files.sort()

    # Add each file's content
    for file_path in html_files:
        # Get section name from directory name
        section_name = os.path.dirname(file_path).replace("-", " ").title()
        if not section_name:
            section_name = "Main Exercise"

        # Add section heading
        doc.add_heading(section_name, level=1)

        full_path = os.path.join(base_dir, file_path)
        content = read_file_content(full_path)

        if content:
            # Add HTML content
            doc.add_heading("HTML", level=2)
            add_code_block(doc, content, "HTML")

            # Check for associated CSS file
            css_path = os.path.join(os.path.dirname(full_path), "style.css")
            if os.path.exists(css_path):
                css_content = read_file_content(css_path)
                if css_content:
                    doc.add_heading("CSS", level=2)
                    add_code_block(doc, css_content, "CSS")

    # Save the document
    output_filename = f"exercise_{exercise_num}_documentation.docx"
    doc.save(output_filename)
    print(f"Generated {output_filename}")


def main():
    while True:
        try:
            exercise_num = int(input("Enter exercise number: "))
            if exercise_num < 1:
                print("Please enter a positive number")
                continue
            break
        except ValueError:
            print("Please enter a valid number")

    create_exercise_doc(exercise_num)


if __name__ == "__main__":
    main()
